from __future__ import annotations

import textwrap
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "assets"
TRAIN_DIR = ROOT / "train"
DOCX_PATH = REPORT_DIR / "项目报告.docx"


BRAND = {
    "ink": RGBColor(15, 23, 42),
    "text": RGBColor(30, 41, 59),
    "muted": RGBColor(100, 116, 139),
    "blue": RGBColor(37, 99, 235),
    "cyan": RGBColor(8, 145, 178),
    "emerald": RGBColor(5, 150, 105),
    "amber": RGBColor(180, 83, 9),
    "border": RGBColor(216, 226, 234),
}


def parse_simple_yaml(path: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip() or line.lstrip().startswith("#") or ":" not in line:
            continue
        key, value = line.split(":", 1)
        data[key.strip()] = value.strip()
    return data


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D8E2EA") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", style: str | None = None, *, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 10.5, bold=bold, color=BRAND["text"])


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 20 if level == 1 else 15 if level == 2 else 12.5, bold=True, color=BRAND["ink"])


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, 9, color=BRAND["muted"])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, 9, bold=True, color=BRAND["ink"])
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run_font(r, 8.8, color=BRAND["text"])
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_image(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    if not path.exists():
        add_paragraph(doc, f"缺失图片：{path}", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def figure_path(name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    return ASSET_DIR / name


def docx_safe_image(source: Path, name: str) -> Path:
    path = figure_path(name)
    if not source.exists():
        return source
    with Image.open(source) as image:
        image.convert("RGB").save(path, "PNG")
    return path


def save_metric_dashboard(df: pd.DataFrame) -> Path:
    path = figure_path("training_metric_dashboard.png")
    plt.style.use("seaborn-v0_8-whitegrid")
    fig, axes = plt.subplots(2, 2, figsize=(13, 8), dpi=180)
    x = df["epoch"]

    axes[0, 0].plot(x, df["metrics/mAP50(B)"], label="Box mAP50", color="#2563eb", linewidth=2)
    axes[0, 0].plot(x, df["metrics/mAP50-95(B)"], label="Box mAP50-95", color="#0891b2", linewidth=2)
    axes[0, 0].set_title("Bounding Box Metrics")
    axes[0, 0].legend()

    axes[0, 1].plot(x, df["metrics/mAP50(P)"], label="Pose mAP50", color="#059669", linewidth=2)
    axes[0, 1].plot(x, df["metrics/mAP50-95(P)"], label="Pose mAP50-95", color="#65a30d", linewidth=2)
    axes[0, 1].set_title("Pose Keypoint Metrics")
    axes[0, 1].legend()

    axes[1, 0].plot(x, df["metrics/precision(B)"], label="Box Precision", color="#2563eb", linewidth=2)
    axes[1, 0].plot(x, df["metrics/recall(B)"], label="Box Recall", color="#7c3aed", linewidth=2)
    axes[1, 0].set_title("Box Precision / Recall")
    axes[1, 0].legend()

    axes[1, 1].plot(x, df["metrics/precision(P)"], label="Pose Precision", color="#059669", linewidth=2)
    axes[1, 1].plot(x, df["metrics/recall(P)"], label="Pose Recall", color="#b45309", linewidth=2)
    axes[1, 1].set_title("Pose Precision / Recall")
    axes[1, 1].legend()

    for ax in axes.ravel():
        ax.set_xlabel("Epoch")
        ax.set_ylim(0.75, 1.01)
        ax.grid(True, alpha=0.28)
    fig.suptitle("YOLO Pose Training Metrics Summary", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.95))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_loss_dashboard(df: pd.DataFrame) -> Path:
    path = figure_path("training_loss_dashboard.png")
    plt.style.use("seaborn-v0_8-whitegrid")
    fig, axes = plt.subplots(2, 2, figsize=(13, 8), dpi=180)
    x = df["epoch"]
    pairs = [
        ("train/box_loss", "val/box_loss", "Box Loss", "#2563eb"),
        ("train/pose_loss", "val/pose_loss", "Pose Loss", "#059669"),
        ("train/cls_loss", "val/cls_loss", "Classification Loss", "#b45309"),
        ("train/dfl_loss", "val/dfl_loss", "DFL Loss", "#7c3aed"),
    ]
    for ax, (train_col, val_col, title, color) in zip(axes.ravel(), pairs):
        ax.plot(x, df[train_col], label=train_col, color=color, linewidth=2)
        ax.plot(x, df[val_col], label=val_col, color="#64748b", linewidth=2, linestyle="--")
        ax.set_title(title)
        ax.set_xlabel("Epoch")
        ax.grid(True, alpha=0.28)
        ax.legend()
    fig.suptitle("Training and Validation Loss Curves", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.95))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_bar_summary(final: pd.Series, best_box: pd.Series, best_pose: pd.Series) -> Path:
    path = figure_path("metric_bar_summary.png")
    labels = ["Box P", "Box R", "Box mAP50", "Box mAP50-95", "Pose P", "Pose R", "Pose mAP50", "Pose mAP50-95"]
    values = [
        final["metrics/precision(B)"],
        final["metrics/recall(B)"],
        final["metrics/mAP50(B)"],
        final["metrics/mAP50-95(B)"],
        final["metrics/precision(P)"],
        final["metrics/recall(P)"],
        final["metrics/mAP50(P)"],
        final["metrics/mAP50-95(P)"],
    ]
    colors = ["#2563eb", "#2563eb", "#0891b2", "#0891b2", "#059669", "#059669", "#65a30d", "#65a30d"]
    fig, ax = plt.subplots(figsize=(13, 4.8), dpi=180)
    bars = ax.bar(labels, values, color=colors)
    ax.set_ylim(0.90, 1.005)
    ax.set_ylabel("Score")
    ax.set_title("Final Epoch Metrics (Epoch 400)")
    ax.grid(axis="y", alpha=0.25)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.002, f"{value:.4f}", ha="center", va="bottom", fontsize=8)
    ax.text(0.01, -0.22, f"Best Box mAP50: epoch {int(best_box['epoch'])}, {best_box['metrics/mAP50(B)']:.5f} | "
            f"Best Pose mAP50: epoch {int(best_pose['epoch'])}, {best_pose['metrics/mAP50(P)']:.5f}",
            transform=ax.transAxes, fontsize=9, color="#475569")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_box(ax, xy, wh, title, subtitle, color, face):
    x, y = xy
    w, h = wh
    rect = plt.Rectangle((x, y), w, h, linewidth=1.8, edgecolor=color, facecolor=face)
    ax.add_patch(rect)
    ax.text(x + w / 2, y + h * 0.63, title, ha="center", va="center", fontsize=12, fontweight="bold", color=color)
    ax.text(
        x + w / 2,
        y + h * 0.34,
        textwrap.fill(subtitle, width=28),
        ha="center",
        va="center",
        fontsize=8.2,
        color="#475569",
        linespacing=1.2,
    )


def arrow(ax, start, end, color="#64748b"):
    ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", color=color, linewidth=2))


def save_architecture_diagram() -> Path:
    path = figure_path("system_architecture_diagram.png")
    fig, ax = plt.subplots(figsize=(13, 7), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("System Architecture: Capture -> Inference -> Storage -> Web Console", fontsize=16, fontweight="bold", pad=18)
    draw_box(ax, (0.4, 5.2), (2.2, 1.3), "Inputs", "Camera / RTSP / Video / Image", "#1e40af", "#eff6ff")
    draw_box(ax, (3.1, 5.2), (2.2, 1.3), "Capture Threads", "CameraCaptureThread x N", "#115e59", "#f0fdfa")
    draw_box(ax, (5.8, 5.2), (2.2, 1.3), "Frame Buffer", "LatestFrameBuffer + Event", "#155e75", "#ecfeff")
    draw_box(ax, (8.5, 5.2), (2.2, 1.3), "Inference", "PCDetectionBackend / BPUDetectionBackend", "#9a3412", "#fff7ed")
    draw_box(ax, (11.2, 5.2), (2.2, 1.3), "Results", "DetectionResult[]", "#5b21b6", "#f5f3ff")
    draw_box(ax, (3.1, 2.0), (2.5, 1.3), "SQLite", "vehicle_history table", "#334155", "#f8fafc")
    draw_box(ax, (6.4, 2.0), (2.5, 1.3), "Web API", "/api/events /api/search", "#047857", "#ecfdf5")
    draw_box(ax, (9.7, 2.0), (2.5, 1.3), "Browser UI", "Query / Logs / Route", "#1e40af", "#eff6ff")
    for x1, x2 in [(2.6, 3.1), (5.3, 5.8), (8.0, 8.5), (10.7, 11.2)]:
        arrow(ax, (x1, 5.85), (x2, 5.85))
    arrow(ax, (12.3, 5.2), (4.4, 3.3), "#059669")
    arrow(ax, (5.6, 2.65), (6.4, 2.65), "#059669")
    arrow(ax, (8.9, 2.65), (9.7, 2.65), "#059669")
    ax.text(0.45, 0.65, "Design intent: keep capture realtime, run inference serially, expose state through snapshots, and store durable events in SQLite.",
            fontsize=10, color="#475569")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_thread_diagram() -> Path:
    path = figure_path("thread_communication_diagram.png")
    fig, ax = plt.subplots(figsize=(13, 7), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Threading and Communication Model", fontsize=16, fontweight="bold", pad=18)
    draw_box(ax, (0.5, 5.5), (2.6, 1.4), "Capture Threads", "Read frames, retry streams, overwrite old frame", "#1e40af", "#eff6ff")
    draw_box(ax, (4.0, 5.5), (2.4, 1.4), "LatestFrameBuffer", "frames dict + Lock + Event", "#155e75", "#ecfeff")
    draw_box(ax, (7.3, 5.5), (2.7, 1.4), "Inference Thread", "wait -> get_all -> detect -> annotate", "#9a3412", "#fff7ed")
    draw_box(ax, (10.9, 5.5), (2.4, 1.4), "Runtime Snapshot", "frames/events/latency/errors", "#047857", "#ecfdf5")
    draw_box(ax, (7.3, 2.3), (2.7, 1.4), "SQLite Writes", "record_occurrence()", "#334155", "#f8fafc")
    draw_box(ax, (10.9, 2.3), (2.4, 1.4), "HTTP Threads", "ThreadingHTTPServer requests", "#5b21b6", "#f5f3ff")
    arrow(ax, (3.1, 6.2), (4.0, 6.2))
    arrow(ax, (6.4, 6.2), (7.3, 6.2))
    arrow(ax, (10.0, 6.2), (10.9, 6.2), "#059669")
    arrow(ax, (8.65, 5.5), (8.65, 3.7), "#059669")
    arrow(ax, (12.1, 3.7), (12.1, 5.5), "#64748b")
    ax.text(0.6, 1.0, "Key rule: model runtime is called by one inference thread only; HTTP handlers read copied snapshots or open short-lived SQLite connections.",
            fontsize=10, color="#475569")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_backend_pipeline_diagram() -> Path:
    path = figure_path("backend_pipeline_diagram.png")
    fig, axes = plt.subplots(2, 1, figsize=(13, 7), dpi=180)
    pipelines = [
        ("PC backend", ["BGR frame", "YOLO .pt", "boxes + 4 keypoints", "crop_plate", "PaddleOCR", "clean text", "DetectionResult"], "#2563eb"),
        ("RDK BPU backend", ["BGR frame", "BGR->NV12", "YOLO .bin", "decode + NMS", "crop_plate", "LPRNet .bin", "DetectionResult"], "#059669"),
    ]
    for ax, (title, nodes, color) in zip(axes, pipelines):
        ax.set_xlim(0, len(nodes))
        ax.set_ylim(0, 1)
        ax.axis("off")
        ax.text(0, 0.86, title, fontsize=14, fontweight="bold", color=color)
        for i, node in enumerate(nodes):
            rect = plt.Rectangle((i + 0.05, 0.25), 0.8, 0.35, linewidth=1.4, edgecolor=color, facecolor="#ffffff")
            ax.add_patch(rect)
            ax.text(i + 0.45, 0.425, node, ha="center", va="center", fontsize=8.5, color="#0f172a")
            if i < len(nodes) - 1:
                ax.annotate("", xy=(i + 1.03, 0.425), xytext=(i + 0.86, 0.425),
                            arrowprops=dict(arrowstyle="->", color="#64748b", linewidth=1.6))
    fig.suptitle("Inference Backend Pipelines", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.92))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_effect_diagram() -> Path:
    path = figure_path("implementation_effect_diagram.png")
    fig, ax = plt.subplots(figsize=(13, 7), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Implementation Effect: Web Console, Query Result, Route and Logs", fontsize=16, fontweight="bold", pad=18)

    draw_box(ax, (0.5, 5.4), (2.6, 1.4), "Monitor Grid", "4 camera tiles with annotated frames", "#1e40af", "#eff6ff")
    draw_box(ax, (4.0, 5.4), (2.6, 1.4), "Live Events", "latest runtime events and timestamps", "#047857", "#ecfdf5")
    draw_box(ax, (7.5, 5.4), (2.6, 1.4), "Search Panel", "plate input, state and suggestions", "#9a3412", "#fff7ed")
    draw_box(ax, (10.9, 5.4), (2.6, 1.4), "Location Result", "camera id, last seen time and crop", "#5b21b6", "#f5f3ff")
    draw_box(ax, (2.4, 2.4), (3.1, 1.4), "Route Map", "highlight path from entrance to C1-C4", "#155e75", "#ecfeff")
    draw_box(ax, (6.6, 2.4), (3.1, 1.4), "SQLite History", "vehicle_history keeps durable records", "#334155", "#f8fafc")
    draw_box(ax, (10.6, 2.4), (2.8, 1.4), "Fallback Demo", "asset images and database-only mode", "#b45309", "#fffbeb")

    arrow(ax, (3.1, 6.1), (4.0, 6.1))
    arrow(ax, (6.6, 6.1), (7.5, 6.1))
    arrow(ax, (10.1, 6.1), (10.9, 6.1))
    arrow(ax, (12.2, 5.4), (8.2, 3.8), "#7c3aed")
    arrow(ax, (8.15, 5.4), (8.15, 3.8), "#059669")
    arrow(ax, (6.6, 3.1), (5.5, 3.1), "#059669")
    arrow(ax, (9.7, 3.1), (10.6, 3.1), "#64748b")

    ax.text(
        0.55,
        0.75,
        "User-visible effect: operators enter a plate number, see the latest camera position, inspect the cropped plate image, and follow the highlighted route.",
        fontsize=10,
        color="#475569",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_end_to_end_logic_diagram() -> Path:
    path = figure_path("end_to_end_logic_diagram.png")
    fig, ax = plt.subplots(figsize=(13, 7), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("End-to-End Implementation Logic", fontsize=16, fontweight="bold", pad=18)

    top_nodes = [
        ("1 Source", "camera / RTSP / video / image", "#1e40af", "#eff6ff"),
        ("2 Capture", "CameraCaptureThread reads frames", "#115e59", "#f0fdfa"),
        ("3 Buffer", "LatestFrameBuffer keeps newest frame", "#155e75", "#ecfeff"),
        ("4 Detect", "Detector.detect frame", "#9a3412", "#fff7ed"),
    ]
    bottom_nodes = [
        ("8 UI", "route, logs, crop and camera tiles", "#5b21b6", "#f5f3ff"),
        ("7 API", "/api/events and /api/search", "#047857", "#ecfdf5"),
        ("6 Store", "record_occurrence() to SQLite", "#334155", "#f8fafc"),
        ("5 Result", "DetectionResult box, pts, text, crop", "#b45309", "#fffbeb"),
    ]
    for i, (title, subtitle, color, face) in enumerate(top_nodes):
        draw_box(ax, (0.45 + i * 3.35, 5.4), (2.65, 1.35), title, subtitle, color, face)
        if i < len(top_nodes) - 1:
            arrow(ax, (3.1 + i * 3.35, 6.08), (3.8 + i * 3.35, 6.08))

    for i, (title, subtitle, color, face) in enumerate(bottom_nodes):
        draw_box(ax, (0.45 + i * 3.35, 2.3), (2.65, 1.35), title, subtitle, color, face)
        if i < len(bottom_nodes) - 1:
            arrow(ax, (3.8 + i * 3.35, 2.98), (3.1 + i * 3.35, 2.98))

    arrow(ax, (12.5, 5.4), (12.5, 3.65), "#059669")
    ax.text(
        0.5,
        1.0,
        "The runtime separates realtime capture, model inference, persistence and presentation so each part can be tested and replaced independently.",
        fontsize=10,
        color="#475569",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)


def add_cover(doc: Document, hero_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("地库车辆定位系统项目报告")
    set_run_font(r, 26, bold=True, color=BRAND["ink"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Garage Vehicle Locator System")
    set_run_font(r, 14, color=BRAND["cyan"])

    if hero_path.exists():
        doc.add_paragraph()
        add_image(doc, hero_path, 16.5, "图 1  系统场景示意：地库摄像头、边缘推理与 Web 控制台")

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["报告类型", "工程项目报告 / 技术实现说明 / 训练结果分析"],
            ["报告范围", "背景需求、系统架构、推理后端、线程通信、训练与部署、测试和改进计划"],
            ["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ["代码仓库", str(ROOT)],
        ],
        [4.2, 12.2],
    )
    doc.add_page_break()


def build_report() -> None:
    REPORT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    df = pd.read_csv(TRAIN_DIR / "results.csv")
    args = parse_simple_yaml(TRAIN_DIR / "args.yaml")
    final = df.iloc[-1]
    first = df.iloc[0]
    best_box = df.loc[df["metrics/mAP50(B)"].idxmax()]
    best_pose = df.loc[df["metrics/mAP50(P)"].idxmax()]
    total_hours = float(final["time"]) / 3600.0

    assets = {
        "metrics": save_metric_dashboard(df),
        "losses": save_loss_dashboard(df),
        "bars": save_bar_summary(final, best_box, best_pose),
        "architecture": save_architecture_diagram(),
        "threading": save_thread_diagram(),
        "backend": save_backend_pipeline_diagram(),
        "effect": save_effect_diagram(),
        "logic": save_end_to_end_logic_diagram(),
    }

    hero = ROOT / "docs" / "tutorial-assets" / "garage-ai-system.png"
    deploy_img = ROOT / "docs" / "tutorial-assets" / "pc-rdk-deployment.png"

    doc = Document()
    configure_doc(doc)
    add_cover(doc, hero)

    add_heading(doc, "一、项目背景与建设目标", 1)
    add_paragraph(
        doc,
        "地下停车场通常存在空间结构复杂、摄像头点位分散、车主找车路径不明确、人工检索效率低等问题。"
        "本项目围绕“根据车牌快速定位车辆最后出现位置”这一核心需求，构建了集多路视频接入、车牌关键点检测、车牌字符识别、轨迹写入和 Web 查询于一体的地库车辆定位系统。"
    )
    add_paragraph(
        doc,
        "系统面向两类场景：其一是在 PC 开发机上使用 YOLO pose 与 PaddleOCR 快速验证算法和界面；其二是在 RDK X5 板端使用 BPU `.bin` 模型进行现场部署。"
        "这种双后端设计使算法、页面和数据库逻辑可以共用，同时把硬件相关细节限制在后端封装层。"
    )
    add_table(
        doc,
        ["目标", "说明"],
        [
            ["实时采集", "支持最多 4 路摄像头、视频、图片或 RTSP 输入，使用最新帧覆盖策略保证低延迟。"],
            ["车牌检测", "YOLO pose 模型输出车牌框和四角关键点，便于透视裁切。"],
            ["字符识别", "PC 端使用 PaddleOCR，RDK 端使用 LPRNet `.bin`。"],
            ["轨迹记录", "SQLite 保存车牌号、摄像头编号、时间和车牌裁切图。"],
            ["可视化查询", "Web 控制台展示监控帧、通行日志、查询结果和路径示意。"],
        ],
        [4.0, 12.8],
    )

    add_table(
        doc,
        ["背景问题", "项目响应"],
        [
            ["车主或管理人员无法快速知道车辆最后出现在哪个区域", "以车牌为索引，记录车辆每次被摄像头识别到的位置，查询时返回最后出现摄像头和时间。"],
            ["地下车库点位多、视频源类型不统一", "输入层统一支持摄像头编号、RTSP、视频文件和图片，后续流程按统一帧对象处理。"],
            ["现场部署可能无法依赖高性能 GPU", "保留 PC 后端用于开发验证，同时提供 RDK X5 BPU 后端用于边缘侧部署。"],
            ["只看模型输出不够，运维需要可解释的事件链", "页面同时展示监控帧、通行日志、车牌裁切图、路径高亮和接口状态。"],
            ["长时间运行需要可恢复的数据", "识别事件写入 SQLite，服务重启后仍可通过数据库查询最近记录。"],
        ],
        [5.3, 11.5],
    )

    add_heading(doc, "二、实现效果与交付形态", 1)
    add_image(doc, assets["effect"], 16.5, "图 2  实现效果：Web 控制台、查询结果、路线图和通行日志")
    add_paragraph(
        doc,
        "项目最终交付形态是一个可直接运行的地库车辆定位控制台。操作者打开浏览器或桌面 WebView 后，可以看到 4 路监控画面区域、实时通行日志、车辆快速查询表单、车辆最后出现位置、车牌裁切图以及路线示意。"
        "这类页面效果的重点不是做静态展示，而是把识别链路产生的数据持续映射为可操作的信息：哪辆车、在哪个摄像头、什么时间出现、是否有可核验的车牌截图。"
    )
    add_paragraph(
        doc,
        "在后台推理运行时开启时，页面每 8 秒调用 `/api/events` 刷新监控画面和事件列表；用户输入车牌后调用 `/api/search`，服务端从 SQLite 中按时间和写入顺序返回最新位置。"
        "如果推理运行时未开启，页面仍能进入数据库浏览模式，并使用 `assets/` 中的样例图填充监控区，便于演示和排障。"
    )
    add_table(
        doc,
        ["效果维度", "用户可见结果", "背后实现"],
        [
            ["实时监控", "页面显示 C1-C4 摄像头 tile，画面带检测框、角点和识别文本。", "`/api/events` 返回 `cameraImages`，后台使用 `draw_annotations()` 生成标注帧并编码为 data URL。"],
            ["车辆定位", "输入车牌后显示最后出现摄像头、时间和状态“已定位”。", "`/api/search` 调用 `query_last_location()`，按 `timestamp DESC, id DESC` 取最新记录。"],
            ["路线提示", "地库平面示意图中高亮入口到目标摄像头的路径。", "前端 `showRoute()` 根据 cameraId 设置 SVG path，并激活对应摄像头点位。"],
            ["车牌核验", "查询结果展示车牌裁切图，便于人工确认识别是否可信。", "数据库 `crop_image` 字段保存 JPEG BLOB，接口转为 base64 data URL。"],
            ["事件追踪", "通行日志持续显示最近识别事件。", "运行时保留最近 100 条事件，数据库可查询最近历史记录。"],
            ["演示降级", "无摄像头或不启用模型时仍可打开页面和查看样例。", "`--backend none` 或无运行时快照时，接口回退到数据库记录和 `asset_images()`。"],
        ],
        [3.4, 6.0, 7.3],
    )
    add_heading(doc, "2.1 页面交互流程", 2)
    add_table(
        doc,
        ["交互步骤", "页面表现", "接口和状态变化"],
        [
            ["打开首页", "显示查询卡片、路线图、摄像头网格和事件列表。", "静态文件由 `GarageWebHandler.serve_file()` 提供；`loadEvents()` 立即拉取 `/api/events`。"],
            ["后台刷新", "摄像头画面和日志周期更新。", "前端 `setInterval(loadEvents, 8_000)`；接口返回 `stats.channels`、`stats.records`、`latency` 和 `errors`。"],
            ["输入车牌", "结果状态从“等待查询”变为“查询中”。", "前端将输入值 `trim().toUpperCase()` 后编码到 `/api/search?plate=`。"],
            ["查到车辆", "显示车牌号、C 号摄像头、最后时间、裁切图和路线。", "接口返回 `ok=true` 与 `result`；前端调用 `setCrop()` 和 `showRoute()`。"],
            ["未查到车辆", "显示“未找到”或相近车牌建议。", "接口返回 `ok=false`；服务端 `query_fuzzy()` 最多给出 6 条相近最新记录。"],
        ],
        [3.1, 5.7, 8.0],
    )
    add_image(doc, docx_safe_image(ROOT / "assets" / "test_plate.jpg", "sample_test_plate.png"), 11.5, "图 3  输入样例 1：用于 PC/RDK 识别链路演示的车牌图片")
    add_image(doc, docx_safe_image(ROOT / "assets" / "test_plate2.jpg", "sample_test_plate2.png"), 11.5, "图 4  输入样例 2：用于验证多图片输入和数据库记录写入")

    add_heading(doc, "三、端到端实现逻辑", 1)
    add_image(doc, assets["logic"], 16.5, "图 5  端到端实现逻辑：输入源、采集、推理、入库、API 与前端")
    add_paragraph(
        doc,
        "系统实现逻辑可以概括为“视频帧进入、模型识别、事件落库、接口读取、页面展示”。"
        "每个环节都尽量保持职责单一：采集线程只负责拿到最新帧，推理线程负责调用模型并形成事件，数据库负责持久化，HTTP 请求只读取快照或数据库结果。"
    )
    add_table(
        doc,
        ["步骤", "输入", "核心处理", "输出或状态"],
        [
            ["1. 输入源归一化", "摄像头编号、RTSP URL、图片或视频文件", "`build_default_sources()` 和命令行 `--inputs` 统一转成 source 列表", "最多 4 路 source，空位可跳过"],
            ["2. 采集", "source", "`CameraCaptureThread.run()` 循环读取帧；图片源按 1 秒间隔重复投递，视频读完后循环", "OpenCV BGR frame"],
            ["3. 最新帧缓冲", "camera_id + frame", "`LatestFrameBuffer.put()` 覆盖旧帧并 set Event", "每路只保留最新帧，避免积压"],
            ["4. 推理", "批量取出的最新帧", "`detector.detect(frame)` 调用 PC 或 BPU 后端", "`DetectionResult[]`"],
            ["5. 事件生成", "识别结果、摄像头编号、耗时", "过滤空文本和过短车牌，记录 latency、confidence 和时间", "运行时 events 和 annotated frame"],
            ["6. 数据持久化", "plate_number、camera_id、crop", "`DBManager.record_occurrence()` 编码 JPEG 并写 SQLite", "`vehicle_history` 新增一条记录"],
            ["7. 接口服务", "HTTP GET", "`/api/events` 读运行时快照，`/api/search` 查数据库", "JSON payload"],
            ["8. 前端呈现", "JSON payload", "更新摄像头 tile、日志、查询结果、路线和裁切图", "可被用户理解的定位结果"],
        ],
        [3.0, 4.2, 6.5, 4.2],
    )
    add_heading(doc, "3.1 核心数据对象", 2)
    add_table(
        doc,
        ["对象", "字段或内容", "作用"],
        [
            ["DetectionResult", "`box`、`pts`、`text`、`confidence`、`crop`", "统一 PC 和 RDK 后端输出，使上层逻辑不关心模型运行平台。"],
            ["运行时快照", "`cameraImages`、`events`、`latency`、`errors`、`running`", "HTTP 线程读取复制后的状态，避免直接操作推理线程内部对象。"],
            ["vehicle_history", "`id`、`plate_number`、`camera_id`、`timestamp`、`crop_image`", "保存车辆历史轨迹；`id` 让同一秒内多条记录也能稳定排序。"],
            ["/api/events", "events、cameraImages、stats、latency、errors", "提供首页实时监控所需全部数据。"],
            ["/api/search", "ok、query、result、suggestions", "提供精确查询结果或模糊建议。"],
        ],
        [3.5, 6.4, 6.8],
    )
    add_heading(doc, "3.2 关键实现规则", 2)
    add_table(
        doc,
        ["规则", "原因", "对应代码"],
        [
            ["只处理每路最新帧", "地库监控更看重当前状态，旧帧积压会造成定位延迟。", "`LatestFrameBuffer` 覆盖写入，`get_all()` 复制后清空。"],
            ["推理线程串行调用模型", "减少模型 runtime 并发风险，特别是 BPU runtime 资源更适合集中调用。", "`WebInferenceRuntime._run_inference()` 逐 camera_id 调用 `detect()`。"],
            ["HTTP 只读快照或短连接查库", "避免 Web 请求阻塞采集和推理线程。", "`snapshot()` 复制状态；查询函数使用局部 sqlite3 连接。"],
            ["事件和错误有限保留", "控制内存增长，前端只需要近期状态。", "`events[:100]` 和 `errors[-20:]`。"],
            ["短文本过滤", "降低误识别写入数据库的概率。", "`if not result.text or len(result.text) < 4: continue`。"],
        ],
        [4.0, 6.0, 6.5],
    )

    add_heading(doc, "四、系统总体架构", 1)
    add_image(doc, assets["architecture"], 16.5, "图 6  系统总体架构：采集、推理、存储与 Web 控制台")
    add_paragraph(
        doc,
        "系统以 `web_app.py` 为主要服务入口。HTTP 服务提供静态页面和 `/api/events`、`/api/search` 两个核心接口；"
        "后台运行时创建采集线程和推理线程，将识别结果同步到运行时快照和 SQLite 数据库。"
    )
    add_table(
        doc,
        ["层次", "核心文件", "职责"],
        [
            ["入口层", "web_app.py / webview_app.py", "启动 HTTP 服务、WebView 壳和推理运行时。"],
            ["前端层", "web/index.html / web/app.js / web/styles.css", "展示实时画面、日志、查询结果和路径示意。"],
            ["后端抽象", "utils/inference.py", "统一 PCDetectionBackend 与 BPUDetectionBackend。"],
            ["板端适配", "utils/ultralytics_yolo_pose.py / preprocess.py / postprocess.py / detect_plate_rdk.py", "封装 RDK X5 BPU 推理、NV12 输入、输出解码和 LPRNet 识别。"],
            ["数据层", "utils/db_manager.py", "SQLite 建表、写入出现记录、末次位置查询和模糊查询。"],
        ],
        [3.2, 6.2, 7.0],
    )
    if deploy_img.exists():
        add_image(doc, deploy_img, 16.5, "图 7  PC 开发环境与 RDK X5 部署环境示意")

    add_heading(doc, "五、推理后端设计", 1)
    add_image(doc, assets["backend"], 16.5, "图 8  PC 与 RDK BPU 推理后端处理链")
    add_paragraph(
        doc,
        "推理层的核心设计是统一接口：`VehiclePlateDetector.detect(frame)` 接收一帧 OpenCV BGR 图像，返回 `DetectionResult[]`。"
        "调用方不需要知道当前使用的是 PC 模型还是 BPU 模型，只需要消费检测框、关键点、识别文本、置信度和车牌裁切图。"
    )
    add_table(
        doc,
        ["字段", "含义", "使用位置"],
        [
            ["box", "车牌矩形框 4 点", "用于 `draw_annotations()` 绘制检测框。"],
            ["pts", "车牌四角关键点", "用于 `crop_plate()` 透视变换裁切车牌。"],
            ["text", "清洗后的车牌号", "写入 SQLite，进入事件日志和查询结果。"],
            ["confidence", "OCR / LPRNet 识别置信度", "用于调试识别质量和页面展示。"],
            ["crop", "车牌裁切图", "以 JPEG BLOB 保存到 `vehicle_history.crop_image`。"],
        ],
        [3.0, 6.0, 7.0],
    )
    add_heading(doc, "5.1 PCDetectionBackend", 2)
    add_paragraph(
        doc,
        "PC 后端使用 `ultralytics.YOLO` 加载 `models/yolo11m-pose-carplate.pt`，从 `result.keypoints.xy` 取得车牌四角关键点，"
        "从 `result.boxes.xyxy` 取得检测框，然后通过 `crop_plate()` 生成透视矫正后的车牌图。字符识别阶段兼容 PaddleOCR 2.x 与 3.x："
        "若 OCR 对象提供 `predict()` 则优先使用新接口，否则退回 `ocr(det=False, cls=False)`。"
    )
    add_heading(doc, "5.2 BPUDetectionBackend", 2)
    add_paragraph(
        doc,
        "RDK 后端首先确认 `hbm_runtime` 可导入，再创建 `UltralyticsYOLOPose` 和 `LPRNetRecognizer`。"
        "`UltralyticsYOLOPose` 将 BGR 图像按模型输入尺寸进行 resize/letterbox，并打包为 NV12 输入；BPU forward 后，后处理模块解码 box、score 和 keypoints，并执行 NMS。"
        "裁切出的车牌图再送入 LPRNet `.bin`，通过字符表解码并计算平均置信度。"
    )

    add_heading(doc, "六、线程安排与通信机制", 1)
    add_image(doc, assets["threading"], 16.5, "图 9  后台线程和 HTTP 请求之间的通信模型")
    add_paragraph(
        doc,
        "系统的实时性来自两个约束：采集线程不做推理，推理线程只处理每路最新帧。"
        "`LatestFrameBuffer` 使用 `threading.Lock` 保护帧字典，用 `threading.Event` 通知推理线程有新数据。"
        "推理线程通过 `get_all()` 复制并清空当前批次，使旧帧自然丢弃，避免视频队列积压。"
    )
    add_table(
        doc,
        ["执行单元", "创建位置", "职责", "同步方式"],
        [
            ["采集线程", "WebInferenceRuntime.start()", "每路输入一个线程，读取摄像头、视频、图片或 RTSP。", "写入 LatestFrameBuffer，使用 lock 和 Event。"],
            ["推理线程", "threading.Thread(target=_run_inference)", "等待新帧、串行推理、标注图像、写库、更新快照。", "读取 FrameBuffer，写运行时状态时使用 runtime lock。"],
            ["HTTP 请求线程", "ThreadingHTTPServer", "处理静态文件、/api/events、/api/search。", "通过 snapshot() 读取运行时状态，数据库查询使用短连接。"],
        ],
        [3.0, 5.0, 6.0, 4.2],
    )

    add_heading(doc, "七、模型训练与结果分析", 1)
    add_paragraph(
        doc,
        f"训练输出位于 `train/`。本次训练任务为 `{args.get('task', 'pose')}`，模型基座为 `{args.get('model', 'unknown')}`，"
        f"训练轮数为 {int(final['epoch'])}，图像尺寸 {args.get('imgsz', '640')}，batch 为 {args.get('batch', 'unknown')}，设备配置为 `{args.get('device', 'unknown')}`。"
        f"`results.csv` 共记录 {len(df)} 个 epoch，最终累计训练时间约 {total_hours:.2f} 小时。"
    )
    add_heading(doc, "7.1 训练配置与设计理由", 2)
    add_table(
        doc,
        ["训练项", "配置或数据", "工程意义"],
        [
            ["任务类型", args.get("task", "pose"), "使用 pose 任务同时预测车牌框和四个角点，服务后续透视裁切。"],
            ["基座模型", args.get("model", "yolo11m-pose.pt"), "中等规模模型兼顾检测能力和可部署性，后续可转换为 RDK BPU `.bin`。"],
            ["训练轮数", f"{int(final['epoch'])} epochs", "长轮次训练用于让关键点回归充分稳定，避免只获得粗略检测框。"],
            ["输入尺寸", str(args.get("imgsz", "640")), "640 输入与 BPU 部署模型尺寸一致，降低训练和部署之间的尺度差异。"],
            ["Batch / 设备", f"batch={args.get('batch', 'unknown')}，device={args.get('device', 'unknown')}", "多卡大 batch 提升训练吞吐，适合 400 轮长训练。"],
            ["AMP", str(args.get("amp", "true")), "混合精度降低显存压力，提高训练速度。"],
            ["输出证据", "results.csv、results.png、PR 曲线、混淆矩阵、train/val batch 图片", "报告中的图表均来自 `train/`，可追溯训练过程和最终效果。"],
        ],
        [3.3, 5.4, 8.0],
    )
    add_heading(doc, "7.2 指标表现与工程含义", 2)
    add_table(
        doc,
        ["指标", "最终 epoch", "最佳值 / 说明"],
        [
            ["Box Precision", f"{final['metrics/precision(B)']:.5f}", "检测框精度，反映误检控制能力。"],
            ["Box Recall", f"{final['metrics/recall(B)']:.5f}", "检测框召回，反映漏检控制能力。"],
            ["Box mAP50", f"{final['metrics/mAP50(B)']:.5f}", f"最佳 epoch {int(best_box['epoch'])}: {best_box['metrics/mAP50(B)']:.5f}"],
            ["Box mAP50-95", f"{final['metrics/mAP50-95(B)']:.5f}", "更严格 IoU 阈值下的综合检测质量。"],
            ["Pose Precision", f"{final['metrics/precision(P)']:.5f}", "关键点精度，反映四角点预测稳定性。"],
            ["Pose Recall", f"{final['metrics/recall(P)']:.5f}", "关键点召回，影响车牌裁切成功率。"],
            ["Pose mAP50", f"{final['metrics/mAP50(P)']:.5f}", f"最佳 epoch {int(best_pose['epoch'])}: {best_pose['metrics/mAP50(P)']:.5f}"],
            ["Pose mAP50-95", f"{final['metrics/mAP50-95(P)']:.5f}", "最终关键点 mAP50-95 已接近 0.995。"],
        ],
        [4.0, 4.2, 8.4],
    )
    add_table(
        doc,
        ["观察结论", "数据依据", "对系统效果的影响"],
        [
            ["检测框已经充分收敛", f"最终 Box mAP50={final['metrics/mAP50(B)']:.5f}，最佳 Box mAP50={best_box['metrics/mAP50(B)']:.5f}", "车牌区域能稳定被框出，是后续裁切和识别的前提。"],
            ["关键点质量更接近部署目标", f"最终 Pose mAP50-95={final['metrics/mAP50-95(P)']:.5f}", "四角点越稳定，透视矫正后的车牌越平整，OCR/LPRNet 越容易识别。"],
            ["训练后期仍有小幅提升", f"最佳 Pose mAP50 出现在 epoch {int(best_pose['epoch'])}", "说明长轮次训练对关键点回归仍有收益。"],
            ["验证损失和训练损失存在差距", f"最终 val/cls_loss={final['val/cls_loss']:.5f}，train/cls_loss={final['train/cls_loss']:.5f}", "现场仍需保留裁切图和日志，便于发现反光、遮挡、夜间等域外情况。"],
        ],
        [4.2, 5.8, 6.8],
    )
    add_image(doc, assets["bars"], 16.5, "图 10  最终 epoch 关键指标柱状图")
    add_image(doc, assets["metrics"], 16.5, "图 11  训练过程中的检测框与关键点指标曲线")
    add_image(doc, assets["losses"], 16.5, "图 12  训练与验证损失曲线")
    add_paragraph(
        doc,
        f"从首轮到最终轮，训练 box loss 由 {first['train/box_loss']:.5f} 降至 {final['train/box_loss']:.5f}，"
        f"训练 pose loss 由 {first['train/pose_loss']:.5f} 降至 {final['train/pose_loss']:.5f}。"
        "曲线显示模型在前期快速收敛，后期主要体现为小幅稳定优化。Pose 指标高于检测框 mAP50-95 的稳定度，对后续透视裁切和字符识别具有直接价值。"
    )

    add_heading(doc, "7.3 训练可视化样例", 2)
    image_rows = [
        (TRAIN_DIR / "labels.jpg", "图 13  数据集标签分布和标注质量概览"),
        (TRAIN_DIR / "results.png", "图 14  Ultralytics 原始训练结果总览图"),
        (TRAIN_DIR / "confusion_matrix_normalized.png", "图 15  归一化混淆矩阵"),
        (TRAIN_DIR / "PosePR_curve.png", "图 16  关键点 PR 曲线"),
        (TRAIN_DIR / "BoxPR_curve.png", "图 17  检测框 PR 曲线"),
        (TRAIN_DIR / "val_batch0_pred.jpg", "图 18  验证集预测样例"),
    ]
    for img, caption in image_rows:
        add_image(doc, img, 15.5, caption)

    add_heading(doc, "八、部署方案", 1)
    add_paragraph(
        doc,
        "PC 部署用于开发联调，推荐 Conda 隔离环境；RDK 部署用于现场运行，推荐使用板端系统 Python 以确保 `hbm_runtime` 可用。"
        "Web 控制台本身不依赖复杂前端构建，`web_app.py` 使用 Python 标准库 HTTP 服务提供静态页面和 JSON API。"
    )
    add_table(
        doc,
        ["部署类型", "关键命令", "适用场景"],
        [
            ["PC Web", "python web_app.py --backend pc --inputs assets/test_plate.jpg assets/test_plate2.jpg", "算法联调、页面演示。"],
            ["PC WebView", "python webview_app.py --backend pc", "需要桌面窗口的演示环境。"],
            ["RDK BPU", "python3 web_app.py --backend bpu --host 0.0.0.0 --inputs /dev/video0", "现场部署、边缘计算。"],
            ["数据库浏览", "python web_app.py --backend none", "只查看已有数据库和页面。"],
        ],
        [3.4, 8.2, 5.0],
    )

    add_heading(doc, "九、测试、风险与改进计划", 1)
    add_paragraph(
        doc,
        "项目提供语法编译、入口参数、核心导入和数据库烟测；RDK 端还可使用 `test/test_headless.py` 在无图形界面环境中验证 BPU 后端。"
        "当前风险主要集中在不同 PaddleOCR 版本返回格式差异、RDK runtime 与 `.bin` 模型版本匹配、RTSP 网络抖动、夜间或反光环境下车牌裁切质量下降等方面。"
    )
    add_table(
        doc,
        ["风险", "影响", "应对策略"],
        [
            ["PaddleOCR 版本差异", "PC 后端识别接口可能变化", "保留 `_create_paddle_ocr` 和 `_recognize_plate_text` 兼容逻辑。"],
            ["BPU runtime 不匹配", "RDK 后端无法加载模型", "部署前用 `import hbm_runtime` 和单图 CLI 验证。"],
            ["视频积压", "实时性下降", "继续使用最新帧覆盖策略，不改为无限队列。"],
            ["车牌裁切质量波动", "OCR/LPRNet 文本为空", "保存裁切图，分析角点顺序、曝光、反光和运动模糊。"],
            ["数据库增长", "长期运行文件变大", "后续可增加归档、去重和保留周期策略。"],
        ],
        [4.0, 5.2, 7.2],
    )

    add_heading(doc, "十、技术参考链接", 1)
    links = [
        ("Python 官方文档", "https://docs.python.org/3/"),
        ("Conda 安装文档", "https://docs.conda.io/projects/conda/en/latest/user-guide/install/index.html"),
        ("OpenCV VideoCapture", "https://docs.opencv.org/4.x/d8/dfe/classcv_1_1VideoCapture.html"),
        ("Ultralytics 文档", "https://docs.ultralytics.com/"),
        ("PaddleOCR 安装文档", "https://github.com/PaddlePaddle/PaddleOCR/blob/main/docs/version3.x/installation.en.md"),
        ("D-Robotics RDK Model Zoo", "https://github.com/D-Robotics/rdk_model_zoo"),
        ("Python sqlite3 文档", "https://docs.python.org/3/library/sqlite3.html"),
        ("systemd 项目文档", "https://systemd.io/"),
    ]
    for name, url in links:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}：")
        set_run_font(r, 10.5, bold=True, color=BRAND["text"])
        add_hyperlink(p, url, url)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录：训练配置摘要", 1)
    selected_args = ["task", "model", "epochs", "batch", "imgsz", "device", "workers", "optimizer", "lr0", "lrf", "momentum", "weight_decay", "mosaic", "fliplr", "amp", "save_dir"]
    add_table(doc, ["参数", "值"], [[k, args.get(k, "")] for k in selected_args], [5.0, 11.5])

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
